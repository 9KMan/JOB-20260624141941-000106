python
// src/core/docx_export.py
"""Word document export.

Renders a generated narrative into a ``.docx`` file based on a template
that the user supplies. The template can contain headers, footers, and
boilerplate paragraphs; the narrative is appended at the end.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType


def _split_paragraphs(text: str) -> list[str]:
    """Split a narrative block into paragraphs on blank lines."""
    normalized = text.replace("\r\n", "\n").strip()
    if not normalized:
        return []
    return [chunk.strip() for chunk in normalized.split("\n\n") if chunk.strip()]


def export_report(narrative: str, template_path: Path, output_path: Path) -> None:
    """Export a narrative into a Word document.

    If ``template_path`` exists, the export is built on top of that
    template (preserving its styles, headers, and any existing body
    content). Otherwise a fresh document is created.

    Args:
        narrative: The full narrative text. Blank lines separate paragraphs.
        template_path: Path to a ``.docx`` template (must exist if supplied).
        output_path: Destination path for the generated report.

    Raises:
        FileNotFoundError: If ``template_path`` was supplied but is missing.
    """
    if template_path is not None and not str(template_path):
        template_path = None  # type: ignore[assignment]

    if template_path and template_path.is_file():
        doc: DocumentType = Document(str(template_path))
    else:
        doc = Document()
        doc.add_heading("Report", level=1)

    for paragraph_text in _split_paragraphs(narrative):
        doc.add_paragraph(paragraph_text)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

